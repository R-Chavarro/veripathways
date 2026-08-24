from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path(r"C:\Users\rosaf\OneDrive\documents\Git\veripathways\deliverables")
OUT.mkdir(exist_ok=True)

NAVY = "173A52"
TEAL = "1B7F79"
PALE = "EAF4F3"
LIGHT = "F3F6F8"
MID = "D2DEE5"
INK = "24333D"
MUTED = "5F6F79"
WHITE = "FFFFFF"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=120, start=140, bottom=120, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    total = sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths[i]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[i] / 1440)
            set_cell_margins(cell)


def mark_first_row_header(table):
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    flag = OxmlElement("w:tblHeader")
    flag.set(qn("w:val"), "true")
    tr_pr.append(flag)


def set_run_font(run, name="Aptos", size=10.5, color=INK, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("VERIPATHWAYS  |  ")
    set_run_font(run, size=8.5, color=MUTED, bold=True)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    fld_run = paragraph.add_run()
    set_run_font(fld_run, size=8.5, color=MUTED)
    fld_run._r.extend([begin, instr, separate, text, end])


def setup_doc():
    d = Document()
    s = d.sections[0]
    s.page_width = Inches(8.5)
    s.page_height = Inches(11)
    s.top_margin = Inches(0.7)
    s.bottom_margin = Inches(0.6)
    s.left_margin = Inches(0.85)
    s.right_margin = Inches(0.85)
    s.header_distance = Inches(0.3)
    s.footer_distance = Inches(0.35)

    styles = d.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08
    for name, size, color, before, after in (
        ("Heading 1", 16, NAVY, 14, 5),
        ("Heading 2", 12.5, TEAL, 10, 3),
        ("Heading 3", 10.5, NAVY, 7, 2),
    ):
        st = styles[name]
        st.font.name = "Aptos Display"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True
    for list_name in ("List Bullet", "List Number"):
        st = styles[list_name]
        st.font.name = "Aptos"
        st.font.size = Pt(10.2)
        st.font.color.rgb = RGBColor.from_string(INK)
        st.paragraph_format.left_indent = Inches(0.25)
        st.paragraph_format.first_line_indent = Inches(-0.18)
        st.paragraph_format.space_after = Pt(3)

    footer = s.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    add_page_field(p)
    return d


def title_block(d, package_label, title, subtitle):
    p = d.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("VERIPATHWAYS")
    set_run_font(r, name="Aptos Display", size=11, color=TEAL, bold=True)
    p = d.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(package_label.upper())
    set_run_font(r, size=9, color=MUTED, bold=True)
    p = d.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(title)
    set_run_font(r, name="Aptos Display", size=24, color=NAVY, bold=True)
    p = d.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(subtitle)
    set_run_font(r, size=11.5, color=MUTED, italic=True)
    callout = d.add_table(rows=1, cols=1)
    callout.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = callout.cell(0, 0)
    shade(cell, PALE)
    set_table_geometry(callout, [9360])
    mark_first_row_header(callout)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Guiding your journey through higher education with confidence, purpose, and a long-term perspective.")
    set_run_font(r, size=10.5, color=NAVY, bold=True)


def section(d, heading, intro=None):
    d.add_heading(heading, level=1)
    if intro:
        d.add_paragraph(intro)


def bullets(d, items):
    for item in items:
        d.add_paragraph(item, style="List Bullet")


def numbered(d, items):
    for item in items:
        d.add_paragraph(item, style="List Number")


def key_value_table(d, rows, compact=False):
    t = d.add_table(rows=0, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (label, value) in enumerate(rows):
        cells = t.add_row().cells
        cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade(cells[0], NAVY if i == 0 else LIGHT)
        shade(cells[1], PALE if i == 0 else WHITE)
        p = cells[0].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label)
        set_run_font(r, size=9.0 if compact else 9.5, color=WHITE if i == 0 else NAVY, bold=True)
        p = cells[1].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        set_run_font(r, size=9.0 if compact else 9.5, color=INK, bold=(i == 0))
    set_table_geometry(t, [2550, 6810])
    if compact:
        for row in t.rows:
            for cell in row.cells:
                set_cell_margins(cell, top=65, start=140, bottom=65, end=140)
    mark_first_row_header(t)
    return t


def disclaimer(d):
    d.add_heading("Important boundaries", level=1)
    p = d.add_paragraph()
    p.add_run("Student authorship. ").bold = True
    p.add_run("The student remains responsible for all application content and final revisions. VeriPathways reviews and comments on materials but does not write essays or application responses on the student's behalf.")
    p = d.add_paragraph()
    p.add_run("Admissions outcomes. ").bold = True
    p.add_run("VeriPathways does not promise, predict, or guarantee admission to any college or university. Admissions decisions are made solely by the institutions to which the student applies.")
    p = d.add_paragraph()
    p.add_run("Financial considerations. ").bold = True
    p.add_run("Any discussion of college cost is intended to help families consider affordability as part of the college decision. It is not financial, tax, or investment advice.")


def next_steps(d):
    d.add_heading("Scope to confirm before enrollment", level=1)
    key_value_table(d, [
        ("SERVICE DETAIL", "CURRENT STATUS"),
        ("Price and duration / active period", "To be confirmed"),
        ("Applications and supplements included", "To be confirmed"),
        ("Review cycles and feedback format", "To be confirmed"),
        ("Turnaround and communication access", "To be confirmed"),
    ], compact=True)


def package_one():
    d = setup_doc()
    title_block(d, "Package 1", "College Conversation + Application Review", "Begin with the bigger college questions. Apply with a clearer sense of purpose.")
    section(d, "Who this is for")
    d.add_paragraph("Designed primarily for students in 10th and 11th grade who want to begin the college conversation before they are deep into the application process. The student and family have time to clarify what matters, examine possible colleges through that lens, and prepare for the application stage without chasing a promised admissions outcome.")
    key_value_table(d, [
        ("BEST FIT", "EARLY COLLEGE EXPLORATION + LATER APPLICATION REVIEW"),
        ("Typical stage", "Primarily grades 10-11; other grades may be considered based on need"),
        ("Primary focus", "Goals, college experience, college list, practical considerations, and application profile"),
        ("Student's role", "Own the choices, write the application, and decide how to use feedback"),
    ])
    section(d, "What the package includes", "Four connected areas of support keep the student at the center of the process.")
    d.add_heading("1. College goals conversation", level=2)
    bullets(d, [
        "Academic interests, possible majors, and early career goals",
        "What the student hopes to learn, experience, and explore during four years of college",
        "Desired challenge, faculty access, mentorship, research, internships, service, entrepreneurship, or clinical exposure",
        "Campus size, location, setting, culture, residential experience, and quality-of-life priorities",
    ])
    d.add_heading("2. College list review", level=2)
    bullets(d, [
        "Whether the list reflects what the student says they want",
        "Meaningful differences among institutions and admissions-selectivity balance",
        "The influence of reputation, marketing, outside pressure, and personal fit",
        "Whether the student would realistically want to attend each school if admitted",
    ])
    d.add_heading("3. Cost and family financial conversation", level=2)
    bullets(d, [
        "Approximate cost of attendance and the family's realistic budget",
        "Alignment between the student and family about willingness to pay",
        "Whether each option remains realistic if aid or scholarships are limited",
    ])
    d.add_heading("4. Application document review", level=2)
    d.add_paragraph("At the active application stage, VeriPathways reviews completed or near-completed materials as one profile and provides professional feedback on how clearly and cohesively the student's experiences, interests, goals, and perspective come through.")
    bullets(d, [
        "Core application or equivalent platform",
        "Activities, honors, and awards sections",
        "Personal statement and a defined number of supplemental essays",
        "Additional information and other agreed-upon materials",
    ])
    section(d, "What the review looks for")
    bullets(d, [
        "The overall impression created by the complete application",
        "Clarity, coherence, authenticity, and context across sections",
        "Whether the student's interests, values, experiences, and intended message are understandable",
        "Questions, gaps, or inconsistencies a reader may notice",
        "Opportunities for the student to clarify or strengthen their own presentation",
    ])
    section(d, "Client journey")
    numbered(d, [
        "Intake - Student and family share background, interests, and initial school ideas.",
        "College conversation - The student explores goals, priorities, and the desired college experience.",
        "College list review - VeriPathways examines the list in the context of those priorities.",
        "Cost conversation - The family considers affordability and willingness to pay.",
        "Application stage - The student prepares their own materials.",
        "Application review and feedback - VeriPathways reviews the profile; the student makes final revisions.",
    ])
    disclaimer(d)
    next_steps(d)
    return d


def package_two():
    d = setup_doc()
    title_block(d, "Package 2", "Application Document Review", "Professional perspective on how your completed application presents you as a whole person.")
    section(d, "Who this is for")
    d.add_paragraph("Designed for students, primarily in 11th and 12th grade, who are actively preparing college applications and do not need the broader college-goals, college-list, or affordability conversation.")
    key_value_table(d, [
        ("BEST FIT", "ACTIVE APPLICATION REVIEW"),
        ("Typical stage", "Primarily grades 11-12 with completed or near-completed materials"),
        ("Primary focus", "How the application materials work together as one student profile"),
        ("Student's role", "Write all content, make final revisions, and retain ownership of the application"),
    ])
    section(d, "What may be reviewed")
    bullets(d, [
        "Common Application or another agreed-upon core application",
        "Activities section",
        "Honors and awards",
        "Personal statement",
        "A defined number of supplemental essays or school-specific essay sets",
        "Additional information section, when applicable",
        "Other materials agreed upon before the review begins",
    ])
    section(d, "The VeriPathways review lens")
    d.add_paragraph("VeriPathways reads the application from the perspective of an experienced higher-education professional. Feedback focuses on what the materials collectively communicate - not on a formula for admission.")
    bullets(d, [
        "Clarity of the student's message and overall impression",
        "Coherence across application sections",
        "Whether experiences, interests, values, and goals are understandable in context",
        "Whether the application feels authentic to the student",
        "Areas that may be confusing, underexplained, or inconsistent",
        "Questions a reader might naturally have",
        "Opportunities for the student to clarify or strengthen their own presentation",
    ])
    section(d, "What this package does not include")
    bullets(d, [
        "College-list review, college-fit advising, or career-goal exploration",
        "Financial or affordability discussion",
        "Application rewriting or essay ghostwriting",
        "Promises, predictions, or guarantees of admission",
        "A formula for admission to a specific institution",
    ])
    section(d, "Client journey")
    numbered(d, [
        "Intake - The student submits the agreed-upon completed or near-completed materials.",
        "Application document review - VeriPathways reads the materials as a complete profile.",
        "Feedback - Strengths, uncertainties, questions, and opportunities for clearer communication are identified.",
        "Student revision - The student decides how to incorporate feedback and makes their own revisions.",
        "Optional follow-up - If included in the final scope, the student discusses feedback or submits a defined second review.",
    ])
    disclaimer(d)
    next_steps(d)
    return d


def save(doc, filename, title):
    props = doc.core_properties
    props.title = title
    props.subject = "VeriPathways college application services"
    props.author = "VeriPathways"
    props.keywords = "VeriPathways, college applications, student advising"
    props.comments = "Client-facing working draft based on the approved service framework."
    doc.save(OUT / filename)


save(package_one(), "VeriPathways_College_Conversation_and_Application_Review.docx", "College Conversation + Application Review")
save(package_two(), "VeriPathways_Application_Document_Review.docx", "Application Document Review")
print(OUT)
